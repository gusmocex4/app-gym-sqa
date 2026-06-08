from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "/Users/unidentified/Downloads/app-gym-main/Plan_de_Pruebas_FitZone.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_base(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Title", 18, RGBColor(46, 116, 181), 0, 12),
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = style_name != "Title"
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def add_title_block(document):
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Plan de Pruebas y Casos de Prueba")
    run.bold = True

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.add_run("Proyecto FitZone / app-gym-main").italic = True


def add_section(document, heading, paragraphs):
    document.add_paragraph(heading, style="Heading 1")
    for item in paragraphs:
        if isinstance(item, tuple) and item[0] == "bullet":
            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(item[1])
        elif isinstance(item, tuple) and item[0] == "number":
            p = document.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(item[1])
        else:
            document.add_paragraph(item)


def add_cases_table(document):
    document.add_paragraph("Casos de Prueba", style="Heading 1")
    cols = ["ID", "Nombre", "Objetivo", "Técnica", "Precondición", "Pasos", "Datos", "Resultado esperado"]
    rows = [
        ["CP-01", "Login vacío", "Validar campos obligatorios", "Partición de equivalencia", "Aplicación disponible", "1. Ir a /login\n2. Enviar formulario vacío", "email vacío\npassword vacía", "Se muestra alerta de campos obligatorios y no se inicia sesión."],
        ["CP-02", "Usuario inexistente", "Rechazar usuario no registrado", "Partición de equivalencia", "Aplicación disponible", "1. Ir a /login\n2. Ingresar credenciales\n3. Enviar", "noexiste@test.com\n12345678", "Se rechaza el acceso y se muestra mensaje de error."],
        ["CP-03", "Password incorrecta", "Rechazar contraseña inválida", "Tabla de decisión", "Usuario existente", "1. Ir a /login\n2. Ingresar email válido\n3. Ingresar clave errónea\n4. Enviar", "user@example.com\nclaveerrada", "Se rechaza el acceso y se informa error de contraseña."],
        ["CP-04", "Login usuario válido", "Permitir acceso a usuario normal", "Tabla de decisión", "Usuario de prueba cargado", "1. Ir a /login\n2. Ingresar credenciales válidas\n3. Enviar", "user@example.com\n123456", "Inicia sesión y redirige a /inicio-user."],
        ["CP-05", "Login admin válido", "Permitir acceso a administrador", "Tabla de decisión", "Admin de prueba cargado", "1. Ir a /login\n2. Ingresar credenciales válidas\n3. Enviar", "admin@example.com\n123456", "Inicia sesión y redirige a /inicio-admin."],
        ["CP-06", "Registro vacío", "Validar campos obligatorios del registro", "Partición de equivalencia", "Aplicación disponible", "1. Ir a /crear-cuenta\n2. Enviar sin datos", "Todos los campos vacíos", "Se muestra alerta de campos obligatorios."],
        ["CP-07", "Contraseña corta", "Validar longitud mínima", "Valor límite", "Aplicación disponible", "1. Ir a /crear-cuenta\n2. Completar formulario\n3. Enviar", "password de 7 caracteres", "Se rechaza el registro e informa mínimo de 8 caracteres."],
        ["CP-08", "Correo duplicado", "Evitar duplicidad de usuario", "Partición de equivalencia", "Existe usuario previo", "1. Ir a /crear-cuenta\n2. Ingresar correo existente\n3. Enviar", "user@example.com", "Se rechaza el registro e informa que el usuario ya existe."],
        ["CP-09", "Registro válido", "Permitir creación de cuenta", "Partición de equivalencia", "Aplicación disponible", "1. Ir a /crear-cuenta\n2. Completar datos válidos\n3. Enviar", "Nombre, apellido, correo nuevo, password >= 8", "Se crea la cuenta y redirige a /login."],
        ["CP-10", "Logout", "Cerrar sesión correctamente", "Prueba basada en estado", "Sesión iniciada", "1. Iniciar sesión\n2. Ir a /logout", "Usuario válido", "La sesión se destruye y redirige a /."],
        ["CP-11", "Persistencia de sesión", "Verificar continuidad tras login", "Prueba basada en estado", "Credenciales válidas", "1. Iniciar sesión\n2. Navegar a pantalla inicial", "Usuario/admin válido", "La sesión permanece activa durante la navegación esperada."],
        ["CP-12", "Rutas principales", "Verificar disponibilidad funcional básica", "Tabla de decisión", "Aplicación disponible", "1. Acceder a rutas principales", "/\n/login\n/crear-cuenta", "Las rutas responden sin fallo funcional."],
    ]

    table = document.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    widths = [Inches(0.6), Inches(1.0), Inches(1.15), Inches(1.0), Inches(1.0), Inches(1.2), Inches(0.9), Inches(1.65)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width

    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, title in enumerate(cols):
        cell = hdr.cells[idx]
        cell.text = title
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.text = value
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                if idx in (0, 1, 3):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_traceability_table(document):
    document.add_paragraph("Matriz de trazabilidad", style="Heading 1")
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    headers = ["Función", "Código relacionado", "Casos"]
    widths = [Inches(1.4), Inches(3.8), Inches(1.3)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, title in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = title
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
    data = [
        ["Login", "controllers/LoginController.php", "CP-01 a CP-05"],
        ["Registro", "controllers/LoginController.php y models/Usuario.php", "CP-06 a CP-09"],
        ["Logout", "controllers/LoginController.php", "CP-10"],
        ["Redirección por rol", "controllers/LoginController.php", "CP-04 y CP-05"],
    ]
    for row_values in data:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.text = value
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)


def build_document():
    document = Document()
    style_base(document)
    add_title_block(document)

    add_section(document, "1. Identificación", [
        "Proyecto: FitZone / app-gym-main",
        "Módulo evaluado: autenticación de usuarios",
        "Característica de calidad: Adecuación funcional",
        "Subcaracterística principal: Corrección funcional",
        "Normas de referencia: ISO/IEC 25010:2023, ISO/IEC/IEEE 29119-4, IEEE 730-2014, IEEE 1012-2024, IEEE 1061-1998",
    ])

    add_section(document, "2. Objetivo", [
        "Verificar que el módulo de autenticación implemente correctamente las funciones de inicio de sesión, registro, cierre de sesión y redirección por rol, produciendo resultados correctos ante entradas válidas e inválidas."
    ])

    add_section(document, "3. Alcance", [
        "Incluye las rutas /login, /crear-cuenta y /logout, así como la redirección a /inicio-admin y /inicio-user.",
        "Se excluyen la recuperación de contraseña, la evaluación de seguridad profunda y las pruebas de rendimiento y usabilidad."
    ])

    add_section(document, "4. Base de diseño", [
        ("bullet", "ISO/IEC 25010:2023: selección de la característica Adecuación funcional y la subcaracterística Corrección funcional."),
        ("bullet", "ISO/IEC/IEEE 29119-4: aplicación de partición de equivalencia, valores límite y tabla de decisión."),
        ("bullet", "IEEE 730-2014: estructura del plan de pruebas."),
        ("bullet", "IEEE 1012-2024: distinción entre verificación y validación."),
        ("bullet", "IEEE 1061-1998: definición de métricas de cobertura y efectividad."),
    ])

    add_section(document, "5. Ítems bajo prueba", [
        ("bullet", "public/index.php"),
        ("bullet", "controllers/LoginController.php"),
        ("bullet", "models/Usuario.php"),
    ])

    add_section(document, "6. Estrategia de prueba", [
        ("bullet", "Pruebas funcionales dinámicas sobre formularios y rutas."),
        ("bullet", "Ejecución manual o automatizada según disponibilidad del entorno."),
        ("bullet", "Validación de mensajes, redirecciones y comportamiento esperado."),
    ])

    add_section(document, "7. Criterios de entrada", [
        ("bullet", "Aplicación desplegada y accesible."),
        ("bullet", "Base de datos disponible."),
        ("bullet", "Usuarios de prueba cargados."),
        ("bullet", "Navegador o entorno de automatización listo."),
    ])

    add_section(document, "8. Criterios de salida", [
        ("bullet", "Todos los casos críticos ejecutados."),
        ("bullet", "Resultados documentados."),
        ("bullet", "Defectos funcionales registrados."),
    ])

    add_section(document, "9. Criterios de aceptación", [
        ("bullet", "El sistema acepta credenciales válidas."),
        ("bullet", "El sistema rechaza entradas inválidas."),
        ("bullet", "El registro funciona con datos válidos."),
        ("bullet", "La contraseña corta se rechaza."),
        ("bullet", "La redirección por rol ocurre correctamente."),
        ("bullet", "Logout destruye la sesión y redirige correctamente."),
    ])

    add_section(document, "10. Riesgos", [
        ("bullet", "Dependencia de conexión a la base de datos."),
        ("bullet", "Los datos de prueba pueden variar entre ejecuciones."),
        ("bullet", "No existe suite automatizada previa en el repositorio."),
    ])

    add_section(document, "11. Métricas", [
        ("bullet", "Cobertura de casos = casos ejecutados / casos planificados."),
        ("bullet", "Tasa de aprobación = casos aprobados / casos ejecutados."),
        ("bullet", "Densidad de defectos = defectos encontrados / casos ejecutados."),
    ])

    add_cases_table(document)
    add_traceability_table(document)

    add_section(document, "12. Conclusión", [
        "El presente plan de pruebas permite evaluar la corrección funcional del módulo de autenticación de FitZone mediante casos observables, repetibles y alineados con normas reconocidas de calidad, prueba y verificación de software."
    ])

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
